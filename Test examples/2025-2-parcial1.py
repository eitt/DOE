"""
parcial_cases (compact PDF version)
==================================
- Entrada de 1 o 2 códigos de alumno.
- Selección de caso por el último dígito del PRIMER código.
- Generación de PDF **compacto** (2 páginas) con encabezado institucional:
    "Primer parcial de Diseño de Experimentos 2025-2"
    "Leonardo H. Talero-Sarmiento, Ph.D."
    "Ingeniería Industrial, Escuela de Estudios Industriales y Empresariales, Universidad Industrial de Santander"
- DataFrame en formato largo (group, value).
- Rúbrica total 5.0 pts (Demostración 0.5; EDA 1.0; ANOVA 1.25; Post Hoc 0.75; Normalidad 0.5; Sensibilidad 0.5; Chequeo de diseño 0.5).

Funciones clave
---------------
- interactive_run(output_dir="/content", compact=True)
- interactive_run_with_docx(output_dir="/content", compact=True, word_template_url=None)
- generate_case_pdf(case_id, user_codes, output_dir="/content", compact=True, course_header_lines=None)
- generate_case_pdf_and_docx(...)
- generate_word_template(...)
"""

from __future__ import annotations
from typing import Dict, List, Tuple, Any, Optional
import os
import textwrap
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages

# Import opcional para Word
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:
    Document = None  # para detectar si no está instalado

__all__ = [
    "ask_user_codes",
    "select_case_by_user_codes",
    "load_case_df",
    "generate_case_pdf",
    "generate_case_pdf_from_user_codes",
    "generate_case_pdf_and_docx",
    "interactive_run",
    "interactive_run_with_docx",
    "build_rubric",
    "generate_word_template",
    "CASES",
]

# -------------------------------
# 1) Diccionario de 10 mini-casos
# -------------------------------

CASES: Dict[int, Dict[str, Any]] = {
    0: {
        'title': "Comparación de resistencia a la tracción por proveedor",
        'scenario': "A quality control manager at a manufacturing plant wants to compare the tensile strength of a new composite material produced by four different suppliers (A, B, C, and D). Due to production schedules, they could not collect the same number of samples from each supplier.",
        'data': {
            'Supplier A': [450, 465, 440, 455, 470],
            'Supplier B': [480, 475, 485, 490],
            'Supplier C': [430, 425, 435],
            'Supplier D': [460, 465, 455, 470, 460, 450]
        },
        'tasks': [
            'EDA: Create box plots for each supplier. Calculate the mean and standard deviation for each group.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant (p < 0.05), use a Tukey's HSD test.",
            'Normality Check: Use a Q-Q plot and a Shapiro-Wilk test on the residuals.',
            'Sensibility Analysis: Re-run the ANOVA and post hoc after changing the highest value for Supplier B (490) to 480. Compare the results.'
        ]
    },
    1: {
        'title': "Productividad por línea de ensamblaje",
        'scenario': "A factory manager wants to compare the average number of units assembled per hour across four different assembly lines (Line 1, Line 2, Line 3, and Line 4). The data were collected on random days, resulting in unequal sample sizes.",
        'data': {
            'Line 1': [25, 28, 26, 30, 29],
            'Line 2': [20, 22, 21, 19],
            'Line 3': [27, 26, 28, 25, 29, 30],
            'Line 4': [23, 24, 25]
        },
        'tasks': [
            'EDA: Create box plots for each line. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after adding a new outlier value of 15 to Line 2. How does this affect the conclusions?'
        ]
    },
    2: {
        'title': "Rendimiento químico por reactor",
        'scenario': "A chemical engineer is investigating the percentage yield of a specific chemical from three different reactors (Reactor A, Reactor B, and Reactor C). Due to maintenance schedules, the number of batches sampled from each reactor is not uniform.",
        'data': {
            'Reactor A': [92.5, 93.1, 91.8, 92.9],
            'Reactor B': [95.0, 94.8, 95.5, 96.1, 95.3],
            'Reactor C': [90.1, 90.5, 89.8, 90.3, 91.0, 89.5]
        },
        'tasks': [
            'EDA: Create box plots for each reactor. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after changing the lowest value in Reactor C (89.5) to 92.0. How does this affect the results?'
        ]
    },
    3: {
        'title': "Desgaste de herramientas por marca",
        'scenario': "An industrial maintenance team is comparing the amount of wear (in micrometers) on cutting tools from three different brands (Brand X, Brand Y, and Brand Z) after a fixed number of operational cycles. The data were collected over several shifts.",
        'data': {
            'Brand X': [15, 17, 16, 18, 14],
            'Brand Y': [12, 11, 13, 10],
            'Brand Z': [20, 21, 19, 22, 23, 20]
        },
        'tasks': [
            'EDA: Create box plots for each brand. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after removing the highest value from Brand Z (23). What is the impact on the significance of the ANOVA?'
        ]
    },
    4: {
        'title': "Efectividad de programas de entrenamiento",
        'scenario': "A human resources department wants to evaluate the effectiveness of four different training programs (Program 1, 2, 3, and 4) for new employees. They measure the time (in minutes) it takes for a new hire to complete a standardized task.",
        'data': {
            'Program 1': [35, 32, 36, 34, 33],
            'Program 2': [40, 38, 41],
            'Program 3': [30, 31, 29, 32, 30, 31],
            'Program 4': [38, 37, 39, 40]
        },
        'tasks': [
            'EDA: Create box plots for each program. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after increasing all values in Program 3 by 5 minutes. Does this change which programs are considered effective?'
        ]
    },
    5: {
        'title': "Resistencia de soldaduras por robot",
        'scenario': "A quality control engineer is assessing the breaking strength (in N/mm²) of welds made by three different welding robots (Robot A, B, and C). The number of welds tested varies due to the robots' availability.",
        'data': {
            'Robot A': [250, 255, 248, 252, 258],
            'Robot B': [240, 245, 242],
            'Robot C': [260, 265, 262, 268, 261, 264]
        },
        'tasks': [
            'EDA: Create box plots for each robot. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after changing the lowest value for Robot B (240) to 250. How does this affect the significance of the pairwise comparisons?'
        ]
    },
    6: {
        'title': "Consumo energético por tipo de motor",
        'scenario': "An energy auditor is comparing the energy consumption (in kWh) of five different types of industrial motors (Motor 1, 2, 3, 4, and 5) over a 24-hour period. Data were collected on different days.",
        'data': {
            'Motor 1': [105, 108, 103, 106],
            'Motor 2': [115, 118, 114],
            'Motor 3': [100, 102, 99, 101, 98],
            'Motor 4': [110, 109, 112],
            'Motor 5': [107, 106, 108, 105]
        },
        'tasks': [
            'EDA: Create box plots for each motor. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after reducing all values for Motor 2 by 5 kWh. What is the impact on the conclusions about energy consumption?'
        ]
    },
    7: {
        'title': "Tiempos de entrega por transportista",
        'scenario': "A logistics manager wants to compare the average delivery time (in days) from four different shipping companies (Company P, Q, R, and S) to a major warehouse.",
        'data': {
            'Company P': [5, 6, 7, 5, 6],
            'Company Q': [4, 3, 5, 4],
            'Company R': [8, 7, 9, 8, 7, 9],
            'Company S': [6, 5, 7]
        },
        'tasks': [
            'EDA: Create box plots for each company. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after changing a value in Company Q (3) to 7. How does this outlier affect the results?'
        ]
    },
    8: {
        'title': "Defectos por línea de producción",
        'scenario': "A quality assurance team is comparing the number of defects found per 1,000 units on three different production lines (Line X, Y, and Z). The number of units inspected varies by line.",
        'data': {
            'Line X': [12, 15, 13, 14],
            'Line Y': [18, 17, 19, 16, 20],
            'Line Z': [10, 11, 9, 12, 10, 11]
        },
        'tasks': [
            'EDA: Create box plots for each line. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after changing all values in Line Z to be 2 higher. Does this change which lines are considered to have a low defect rate?'
        ]
    },
    9: {
        'title': "Secado de pintura por condición ambiental",
        'scenario': "A product development team wants to compare the drying time (in minutes) of a new industrial paint formula under three different environmental conditions (Condition 1, 2, and 3). The data were collected at different intervals.",
        'data': {
            'Condition 1': [45, 42, 48, 46],
            'Condition 2': [55, 52, 58, 56, 54],
            'Condition 3': [35, 38, 36, 34, 39, 37]
        },
        'tasks': [
            'EDA: Create box plots for each condition. Calculate descriptive statistics.',
            'ANOVA: Perform a one-way ANOVA. Construct the ANOVA table.',
            "Post Hoc: If ANOVA is significant, use a Tukey's HSD test.",
            'Normality Check: Check the normality of residuals.',
            'Sensibility Analysis: Re-run the analysis after removing the lowest value from Condition 3 (34). Does this change the conclusion about the drying times?'
        ]
    }
}

# ------------------------------------------------
# 2) Utilidades básicas (selección y DataFrame)
# ------------------------------------------------

def case_id_from_user_code(user_code: str) -> int:
    """Devuelve el ID de caso (0-9) según el último dígito del código; si no es dígito, usa 0."""
    if not user_code:
        return 0
    last = user_code.strip()[-1]
    return int(last) % 10 if last.isdigit() else 0

def select_case_by_user_codes(user_codes: List[str]) -> Tuple[int, Dict[str, Any]]:
    """Selecciona el caso usando el PRIMER código de la lista. Retorna (case_id, case_dict)."""
    if not user_codes:
        raise ValueError("Debes proporcionar al menos un código de usuario.")
    cid = case_id_from_user_code(user_codes[0])
    return cid, CASES[cid]

def load_case_df(case_id: int) -> pd.DataFrame:
    """Convierte el bloque 'data' del caso a formato largo (tidy). Columnas: 'group', 'value'."""
    if case_id not in CASES:
        raise KeyError(f"case_id fuera de rango (0-9): {case_id}")
    data = CASES[case_id]['data']
    rows = []
    for group, values in data.items():
        for v in values:
            rows.append({'group': group, 'value': v})
    return pd.DataFrame(rows)

# ---------------------------------
# 3) Entrada interactiva de códigos
# ---------------------------------

def ask_user_codes() -> List[str]:
    """
    Pide por consola si hay 1 o 2 alumnos y solicita sus códigos.
    Retorna una lista con 1 o 2 códigos.
    """
    while True:
        try:
            n = int(input("¿Cuántos alumnos? (1 o 2): ").strip())
            if n in (1, 2):
                break
            print("Por favor ingrese 1 o 2.")
        except ValueError:
            print("Entrada no válida. Intente de nuevo.")
    codes = []
    for i in range(n):
        code = input(f"Código del alumno {i+1}: ").strip()
        while not code:
            print("El código no puede estar vacío.")
            code = input(f"Código del alumno {i+1}: ").strip()
        codes.append(code)
    return codes

# --------------------
# 4) Rúbrica (5.0 pts)
# --------------------

def build_rubric() -> pd.DataFrame:
    """
    Construye la rúbrica solicitada (total 5.0):
      - Demostración: 0.5
      - Desarrollo de ejercicios (4.0) desglosado en:
          EDA 1.0, ANOVA 1.25, Post Hoc 0.75, Normalidad 0.5, Sensibilidad 0.5
      - Chequeo de diseño (factores, réplicas, n, datos perdidos, variable respuesta): 0.5
    """
    items = [
        {"Componente": "Demostración (ortogonalidad)", "Puntos": 0.5, "Criterios": "Claridad, rigor, conexión con el caso."},
        {"Componente": "EDA", "Puntos": 1.0, "Criterios": "Boxplots por grupo, estadísticas descriptivas, interpretación inicial."},
        {"Componente": "ANOVA", "Puntos": 1.25, "Criterios": "Modelo correcto, tabla ANOVA, interpretación de F y p-valor."},
        {"Componente": "Post Hoc", "Puntos": 0.75, "Criterios": "Tukey HSD (si aplica), pares significativos, conclusiones."},
        {"Componente": "Normalidad de residuos", "Puntos": 0.5, "Criterios": "Q-Q plot, aleatoriedad, homocedasticidad y discusión de supuestos."},
        {"Componente": "Análisis de sensibilidad", "Puntos": 0.5, "Criterios": "Cambio propuesto, comparación de resultados y discusión."},
        {"Componente": "Chequeo de diseño", "Puntos": 0.5, "Criterios": "Factores, réplicas, número de datos, datos perdidos, variable respuesta."},
    ]
    return pd.DataFrame(items)

# -----------------------------
# 5) Generación de PDF (compacto)
# -----------------------------

def _wrap_text(s: str, width: int = 110) -> str:
    wrapped = []
    for para in s.splitlines():
        if not para.strip():
            wrapped.append("")
        else:
            wrapped.extend(textwrap.wrap(para, width=width))
    return "\n".join(wrapped)

def _add_text_page(pdf: PdfPages, title: str, body: str, fontsize: int = 10):
    fig = plt.figure(figsize=(8.27, 11.69))  # A4 vertical
    fig.suptitle(title, fontsize=14, y=0.98)
    ax = fig.add_subplot(111)
    ax.axis('off')
    ax.text(
        0.03, 0.97,
        _wrap_text(body, width=110),
        va='top', ha='left', fontsize=fontsize, family='monospace'
    )
    pdf.savefig(fig, bbox_inches='tight')
    plt.close(fig)

def _truncate_values(vals: List[float], max_show: int = 12) -> str:
    if len(vals) <= max_show:
        return str(vals)
    return str(vals[:6] + ["…"] + vals[-2:])

def generate_case_pdf(
    case_id: int,
    user_codes: List[str],
    output_dir: str = "/content",
    compact: bool = True,
    course_header_lines: Optional[List[str]] = None,
) -> str:
    """
    Genera un PDF (compacto en 2 páginas por defecto) con:
    - Encabezado del curso/periodo/docente/institución.
    - Caso (sin revelar título en consola) + escenario.
    - Tareas + Demostración (según 1 o 2 alumnos).
    - Rúbrica + Resumen del dataset + Chequeo de diseño.
    """
    case = CASES[case_id]
    os.makedirs(output_dir, exist_ok=True)
    safe_codes = "_".join([c.replace("/", "-") for c in user_codes])
    fname = f"parcial_case{case_id}_{safe_codes}.pdf"
    out_path = os.path.join(output_dir, fname)

    # Encabezado por defecto
    if course_header_lines is None:
        course_header_lines = [
            "Primer parcial de Diseño de Experimentos 2025-2",
            "Leonardo H. Talero-Sarmiento, Ph.D.",
            "Ingeniería Industrial, Escuela de Estudios Industriales y Empresariales, Universidad Industrial de Santander",
        ]

    header_block = "\n".join(course_header_lines) + f"\nCódigos: {' / '.join(user_codes)}"

    scenario = f"{header_block}\n\nScenario:\n{case['scenario']}"
    # Resumen dataset (compacto)
    data = case['data']
    lines = []
    total_n = 0
    for g, vals in data.items():
        total_n += len(vals)
        show_vals = _truncate_values(vals, max_show=12)
        lines.append(f"  • {g}: n={len(vals)}; valores={show_vals}")
    dataset_summary = f"Resumen del dataset (total n={total_n}):\n" + "\n".join(lines)

    tasks = "Tasks:\n" + "\n".join([f"  - {t}" for t in case['tasks']])

    # Demostración condicional
    if len(user_codes) == 1:
        demo_title = "Demostración (0.5 pts) – Ortogonalidad de los contrastes"
        demo_body = (
            "Defina ortogonalidad de contrastes; pruebe que ⟨c_i, c_j⟩ = 0 ⇒ independencia entre estimadores; "
            "construya ≥ 2 contrastes ortogonales con los grupos del caso; cierre explicando por qué favorece la interpretación."
        )
    else:
        demo_title = "Demostración (0.5 pts) – Ortogonalidad y factores independientes (bloques)"
        demo_body = (
            "Defina ortogonalidad entre subespacios de tratamiento y bloque; muestre separación de efectos y suma de cuadrados aditiva; "
            "ilustre con matriz de diseño y proyecciones; concluya ventajas para inferencia y eficiencia."
        )

    rubric_df = build_rubric()
    rubric_txt = "Rúbrica (Total 5.0 pts):\n" + "\n".join(
        [f"  - {row.Componente}: {row.Puntos} pts. {row.Criterios}" for _, row in rubric_df.iterrows()]
    )

    design_check = (
        "Chequeo de diseño (0.5 pts)\n"
        "  1) ¿Cuántos factores hay? (principal y si existen bloques)\n"
        "  2) ¿Cuántas réplicas por nivel/grupo?\n"
        "  3) ¿Número total de observaciones (n)?\n"
        "  4) ¿Datos perdidos o atípicos sospechosos? Justifique.\n"
        "  5) ¿Variable respuesta y su escala (unidades)?"
    )

    # Páginas PDF (compact = 2 páginas)
    with PdfPages(out_path) as pdf:
        # Página 1: Encabezado + Scenario + Dataset (datos del examen)
        page1 = f"{scenario}\n\n{dataset_summary}"
        _add_text_page(pdf, "Datos del examen", page1, fontsize=10)

        # Página 2: Tareas → Demostración → Preguntas sobre diseño → Rúbrica
        page2 = f"{tasks}\n\n{demo_title}\n{demo_body}\n\n{design_check}\n\n{rubric_txt}"
        _add_text_page(pdf, "Tareas, Demostración, Diseño y Rúbrica", page2, fontsize=10)

    return out_path

# -----------------------------
# 5bis) Plantilla Word (.docx)
# -----------------------------

def _ensure_docx():
    if Document is None:
        raise ImportError(
            "Falta la dependencia 'python-docx'. En Colab ejecútalo primero: !pip -q install python-docx"
        )

def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def _add_para(doc, text, bold=False, italic=False, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    font = run.font
    font.size = Pt(size)
    p_format = p.paragraph_format
    p_format.space_after = Pt(space_after)
    return p

def generate_word_template(
    case_id: int,
    user_codes: List[str],
    output_dir: str = "/content",
    course_header_lines: Optional[List[str]] = None,
    download_template_url: Optional[str] = None,
) -> str:
    """
    Genera una plantilla Word (.docx) para que el/la estudiante responda el parcial.
    - Títulos y subtítulos según tareas (EDA, ANOVA, Post Hoc, Normalidad, Sensibilidad).
    - 'Demostración' se deja indicada 'A mano'.
    - Incluye 'Chequeo de diseño' con las 5 preguntas.
    Si se pasa 'download_template_url', descargará esa plantilla en vez de generarla.
    Devuelve la ruta del .docx guardado en 'output_dir'.
    """
    import urllib.request
    os.makedirs(output_dir, exist_ok=True)
    safe_codes = "_".join([c.replace("/", "-") for c in user_codes])
    out_docx = os.path.join(output_dir, f"plantilla_parcial_case{case_id}_{safe_codes}.docx")

    # Opción: descargar plantilla desde GitHub (si proporcionas URL raw)
    if download_template_url:
        urllib.request.urlretrieve(download_template_url, out_docx)
        return out_docx

    # Generar documento desde cero
    _ensure_docx()
    case = CASES[case_id]

    if course_header_lines is None:
        course_header_lines = [
            "Primer parcial de Diseño de Experimentos 2025-2",
            "Leonardo H. Talero-Sarmiento, Ph.D.",
            "Ingeniería Industrial, Escuela de Estudios Industriales y Empresariales, Universidad Industrial de Santander",
        ]

    doc = Document()

    # Portada/encabezado
    _add_heading(doc, course_header_lines[0], level=1)
    for line in course_header_lines[1:]:
        _add_para(doc, line, size=11)
    _add_para(doc, f"Códigos: {' / '.join(user_codes)}", bold=True, size=11)
    _add_para(doc, f"Caso asignado: {case_id}", italic=True, size=10)

    doc.add_paragraph("")  # espacio

    # 1) Datos del examen
    _add_heading(doc, "1. Datos del examen", level=2)
    _add_para(doc, "Scenario:", bold=True)
    _add_para(doc, case['scenario'], size=11)

    # Resumen del dataset
    data = case['data']
    _add_para(doc, "Resumen del dataset:", bold=True)
    for g, vals in data.items():
        _add_para(doc, f"• {g}: n={len(vals)}; valores={vals}", size=10)

    doc.add_paragraph("")

    # 2) Desarrollo de tareas
    _add_heading(doc, "2. Desarrollo de tareas", level=2)

    _add_heading(doc, "2.1 EDA (1.0 pts)", level=3)
    _add_para(doc, "Incluya boxplots por grupo y estadísticas descriptivas (media, sd). Interprete patrones o outliers.", size=11)
    _add_para(doc, "Respuesta:", bold=True); doc.add_paragraph(" ")

    _add_heading(doc, "2.2 ANOVA (1.25 pts)", level=3)
    _add_para(doc, "Modele ANOVA de 1 factor, reporte tabla ANOVA (F, p-valor) e interprete.", size=11)
    _add_para(doc, "Respuesta:", bold=True); doc.add_paragraph(" ")

    _add_heading(doc, "2.3 Post Hoc (0.75 pts)", level=3)
    _add_para(doc, "Si ANOVA es significativo, haga Tukey HSD; indique pares significativos y conclusión.", size=11)
    _add_para(doc, "Respuesta:", bold=True); doc.add_paragraph(" ")

    _add_heading(doc, "2.4 Normalidad de residuos (0.5 pts)", level=3)
    _add_para(doc, "Incluya Q–Q plot, Shapiro–Wilk y comente supuestos (aleatoriedad, homocedasticidad).", size=11)
    _add_para(doc, "Respuesta:", bold=True); doc.add_paragraph(" ")

    _add_heading(doc, "2.5 Análisis de sensibilidad (0.5 pts)", level=3)
    _add_para(doc, "Aplique el cambio/escenario indicado en el enunciado y compare conclusiones.", size=11)
    _add_para(doc, "Respuesta:", bold=True); doc.add_paragraph(" ")

    doc.add_paragraph("")

    # 3) Demostración
    _add_heading(doc, "3. Demostración (0.5 pts) – A mano", level=2)
    _add_para(doc, "Entregue la demostración escrita a mano y adjúntela como anexo o imagen aparte.", italic=True, size=10)

    doc.add_paragraph("")

    # 4) Chequeo de diseño
    _add_heading(doc, "4. Chequeo de diseño (0.5 pts)", level=2)
    _add_para(doc, "Responda explícitamente:", bold=True)
    _add_para(doc, "1) ¿Cuántos factores hay? (principal y si existen bloques)", size=11); doc.add_paragraph(" ")
    _add_para(doc, "2) ¿Cuántas réplicas por nivel/grupo?", size=11); doc.add_paragraph(" ")
    _add_para(doc, "3) ¿Número total de observaciones (n)?", size=11); doc.add_paragraph(" ")
    _add_para(doc, "4) ¿Datos perdidos o atípicos sospechosos? Justifique.", size=11); doc.add_paragraph(" ")
    _add_para(doc, "5) ¿Variable respuesta y su escala (unidades)?", size=11); doc.add_paragraph(" ")

    # 5) (Opcional) Conclusiones
    _add_heading(doc, "5. Conclusiones generales", level=2)
    doc.add_paragraph(" ")

    # Guardar
    doc.save(out_docx)
    return out_docx

# ---------------------------------------------
# 6) Ejecutores (PDF y/o DOCX)
# ---------------------------------------------

def generate_case_pdf_and_docx(
    case_id: int,
    user_codes: List[str],
    output_dir: str = "/content",
    compact: bool = True,
    course_header_lines: Optional[List[str]] = None,
    word_template_url: Optional[str] = None,  # si quieres bajar plantilla desde GitHub
) -> Tuple[str, str]:
    """
    Crea PDF e inmediatamente una plantilla Word (.docx) en blanco.
    Devuelve (pdf_path, docx_path).
    """
    pdf_path = generate_case_pdf(
        case_id, user_codes, output_dir=output_dir, compact=compact, course_header_lines=course_header_lines
    )
    docx_path = generate_word_template(
        case_id, user_codes, output_dir=output_dir, course_header_lines=course_header_lines, download_template_url=word_template_url
    )
    return pdf_path, docx_path

def interactive_run(output_dir: str = "/content", compact: bool = True) -> Tuple[pd.DataFrame, str, int]:
    """
    Flujo completo sin imprimir en consola el nombre del caso:
      - Pide 1 o 2 códigos.
      - Selecciona el caso por el primer código.
      - Genera el PDF (2 páginas en modo compacto) con encabezado institucional en `output_dir`.
      - Devuelve (df, pdf_path, case_id).
    """
    codes = ask_user_codes()
    case_id, _ = select_case_by_user_codes(codes)
    df = load_case_df(case_id)
    pdf_path = generate_case_pdf(case_id, codes, output_dir=output_dir, compact=compact)
    return df, pdf_path, case_id

def interactive_run_with_docx(
    output_dir: str = "/content",
    compact: bool = True,
    word_template_url: Optional[str] = None,  # RAW de GitHub si usas una plantilla fija
) -> Tuple[pd.DataFrame, str, str, int]:
    """
    Igual a interactive_run, pero además crea la plantilla Word y la devuelve.
    Retorna: (df, pdf_path, docx_path, case_id)
    """
    codes = ask_user_codes()
    case_id, _ = select_case_by_user_codes(codes)
    df = load_case_df(case_id)
    pdf_path, docx_path = generate_case_pdf_and_docx(
        case_id, codes, output_dir=output_dir, compact=compact, word_template_url=word_template_url
    )
    return df, pdf_path, docx_path, case_id

def generate_case_pdf_from_user_codes(
    user_codes: List[str],
    output_dir: str = "/content",
    compact: bool = True,
    course_header_lines: Optional[List[str]] = None,
) -> Tuple[int, str]:
    """Atajo: selecciona case_id por los códigos y crea el PDF. Retorna (case_id, pdf_path)."""
    case_id, _ = select_case_by_user_codes(user_codes)
    pdf_path = generate_case_pdf(case_id, user_codes, output_dir=output_dir, compact=compact, course_header_lines=course_header_lines)
    return case_id, pdf_path
